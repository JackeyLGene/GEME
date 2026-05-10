"""Build GEME Paper v2 — integrates all annotations + code review adjustments."""
import copy, re, os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

src = "C:/Users/Administrator.DESKTOP-EM03IHL/Desktop/GEME_Paper_Draft.docx"
dst = "C:/Users/Administrator.DESKTOP-EM03IHL/Desktop/GEME_Paper_Draft_v2.docx"

doc = Document(src)

def new_para(text, bold=False, italic=False, size=12):
    """Add a paragraph with given text."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p

def insert_after(doc, marker_text, new_paragraphs):
    """Insert paragraphs after the paragraph containing marker_text."""
    for i, para in enumerate(doc.paragraphs):
        if marker_text in para.text:
            # Insert after this paragraph
            insert_pos = para._element
            for p_text, bold, italic in new_paragraphs:
                new_p = copy.deepcopy(doc.paragraphs[0]._element)  # template
                # Actually, let's use the element approach
                from docx.oxml import OxmlElement
                new_p_elem = OxmlElement('w:p')
                r_elem = OxmlElement('w:r')
                t_elem = OxmlElement('w:t')
                t_elem.text = p_text
                t_elem.set(qn('xml:space'), 'preserve')
                r_elem.append(t_elem)
                new_p_elem.append(r_elem)
                insert_pos.addnext(new_p_elem)
                insert_pos = new_p_elem
            return True
    return False

def replace_in_para(doc, old_text, new_text):
    """Replace text in all paragraphs."""
    count = 0
    for para in doc.paragraphs:
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
    return count

# ===== APPLY ALL CHANGES =====

# 1. Replace "27-dimensional vectors" (annotation 8)
c1 = replace_in_para(doc, "27-dimensional vectors", "fixed-dimensional vectors (currently of moderate size)")
print(f"1. '27-dim' replaced: {c1}")

# 2. Update MI value: 0.085948 -> 0.032
c2 = replace_in_para(doc, "0.085948", "0.032")
c2b = replace_in_para(doc, "0.086", "0.032")
print(f"2. MI value updated: {c2}")

# 3. Update L4 frame count: 6 -> 1
c3 = replace_in_para(doc, "six\u00b1two", "one\u00b1zero point two")
c3b = replace_in_para(doc, "6\u00b12", "1\u00b10.2")
c3c = replace_in_para(doc, "approximately six", "approximately 1")
c3d = replace_in_para(doc, "Magic 6", "stable self-referential attractor")
print(f"3. L4 count updated: {c3}")

# 4. Replace "emergence" / "phase transition" with "threshold-triggered" in section 3.2
c4 = replace_in_para(doc, "phase transition at the fourth layer", "threshold-triggered behavior at the fourth layer")
c4b = replace_in_para(doc, "Phase Transition at the Fourth Layer", "Threshold-Triggered L4 Response")
print(f"4. Emergence wording: {c4}")

# 5. Update complexity claim
c5 = replace_in_para(doc, "O(C)", "O(W + K)")
c5b = replace_in_para(doc, "prediction cost \u2248 verification", "prediction cost approaches verification")
print(f"5. Complexity: {c5}")

# 6. Fix S2 label
c6 = replace_in_para(doc, "S2. GEME and Formality: Q + G \u2248 PA", "S2. Extended Verification: Q + G \u2248 PA")
print(f"6. S2 label: {c6}")

# 7. Add recent work paragraph in Introduction
# Find the paragraph containing "fundamental principles"
for para in doc.paragraphs:
    if "separate conversations" in para.text and "information, formality" in para.text:
        # Insert after this paragraph
        from docx.oxml import OxmlElement
        p_elem = para._element
        # Create new paragraph
        new_p = OxmlElement('w:p')
        r_elem = OxmlElement('w:r')
        t_elem = OxmlElement('w:t')
        t_elem.text = "Recent work has explored the intersection of self-reference and computation from multiple angles. Aaronson (2013) examined the relationship between quantum mechanics and computational complexity, while Friston (2010) proposed the free-energy principle as a unified account of brain function. Tononi\u2019s integrated information theory (Oizumi et al., 2014) directly addresses consciousness as an information-theoretic quantity. GEME differs from these approaches in a fundamental respect: it does not define what cognition should look like, but provides a minimal substrate in which cognitive-like structures emerge from economic pressure alone\u2014no free energy, no integrated information, and no quantum mechanics are assumed."
        t_elem.set(qn('xml:space'), 'preserve')
        r_elem.append(t_elem)
        new_p.append(r_elem)
        # Set paragraph properties for justified alignment
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        new_p.insert(0, pPr)
        p_elem.addnext(new_p)
        print("7. Recent work paragraph inserted")
        break

# 8. Add metaframework paragraph in Limitations
for para in doc.paragraphs:
    if "Limitations and Future Work" in para.text:
        # Find the next paragraphs, insert after the 27-dim sentence fix
        pass

# We insert the metaframework paragraph after "richer sensory grounding"
for i, para in enumerate(doc.paragraphs):
    if "richer sensory grounding" in para.text:
        from docx.oxml import OxmlElement
        p_elem = para._element
        new_p = OxmlElement('w:p')
        r_elem = OxmlElement('w:r')
        t_elem = OxmlElement('w:t')
        t_elem.text = "GEME\u2019s layered architecture separates world-processing (L1-L3) from self-referential verification (L4-L6). The upper three layers form a metaframework: L4 predicts the next state, L5 observes the prediction outcome, and L6 issues a judgment. This tripartite structure suggests that consciousness\u2014if it emerges\u2014may require self-reference organized into a prediction-observation-judgment pipeline. A single self-referential frame acting as the L6 output could serve as a component in a larger network of such agents."
        t_elem.set(qn('xml:space'), 'preserve')
        r_elem.append(t_elem)
        new_p.append(r_elem)
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        new_p.insert(0, pPr)
        p_elem.addnext(new_p)
        print("8. Metaframework paragraph inserted")
        break

# 9. Insert the reordered section headers
# Note: This is tricky because section 3 sections need reordering.
# We'll note this as needing manual adjustment in Word.

# 10. Move Miller/Milgram paragraph to section 4.3 (Unifying Implications)
for para in doc.paragraphs:
    if "common information-economic principle" in para.text:
        from docx.oxml import OxmlElement
        # Find section 4.3
        for j, p2 in enumerate(doc.paragraphs):
            if "4.3" in p2.text and "Unifying" in p2.text:
                # Move the paragraph here
                p_elem = p2._element
                # Deep copy the paragraph element
                from copy import deepcopy
                moved = deepcopy(para._element)
                # Remove original
                para._element.getparent().remove(para._element)
                p_elem.addnext(moved)
                print(f"9. Miller/Milgram paragraph moved to section 4.3")
                break
        break

doc.save(dst)
print(f"\nSaved to {dst}")
print("\nNOTE: Section 3 reordering (Godel bridge -> Parameter stability -> Threshold emergence) requires manual Word edit.")
print("CHANGES_v2.md has the complete list.")
