# -*- coding: utf-8 -*-
"""生成中文版论文提纲（含公式）"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

doc = Document()

# --- 页面设置 ---
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# --- 样式 ---
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(3)

for level, size in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)]:
    s = doc.styles[level]
    s.font.name = '黑体'
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

def add_para(text, indent=True):
    p = doc.add_paragraph(text)
    if indent: p.paragraph_format.first_line_indent = Pt(24)
    return p

def add_formula(formula_text):
    """添加公式行（居中）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(formula_text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.italic = True
    return p

def add_axiom(number, name, text):
    p = doc.add_paragraph()
    r = p.add_run(f'公理 {number}（{name}）：')
    r.bold = True
    p.add_run(text)
    p.paragraph_format.first_line_indent = Pt(24)
    return p

# ==================== 标题页 ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('GEME：一个自指棱镜框架')
r.font.size = Pt(22)
r.font.bold = True
r.font.name = '黑体'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('——竞争性帧经济中的涌现结构')
r.font.size = Pt(14)
r.font.name = '黑体'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('刘洁琪')
r.font.size = Pt(14)
r.font.name = '楷体'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('独立研究者')
r.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DOI: 10.5281/zenodo.20059553')
r.font.size = Pt(10)
r.font.name = 'Times New Roman'

doc.add_paragraph()

# ==================== 摘要 ====================
doc.add_heading('摘要', level=1)
add_para(
    "理解认知架构的基元机制是核心挑战。本文提出生成式经济记忆实体（GEME），"
    "一个基于三条不可调整公理的极简计算模型：竞争性帧合并、自适应遗忘和自指观测。"
    "系统无自由参数，仅含三个固定结构常数。在运行过程中，帧经济自发产生三类涌现现象："
    "（1）自指操作在信息论上的资源中性——自指帧与输入信道的互信息趋于零；"
    "（2）第四层自指层的行为相变，对应于元认知监控的涌现；"
    "（3）L4帧数收敛到约6个的稳定吸引子，与米勒7±2、米尔格拉姆六度分割、"
    "及形式系统Q+G≈PA在数值上结构一致。GEME不模拟认知——它提供一个最小计算基元，"
    "认知类结构在其中自发产生。模型完全开源可复现。"
)

# ==================== 1. 引言 ====================
doc.add_heading('1. 引言', level=1)
add_para(
    "将香农的信息论（Shannon, 1948）、哥德尔的不完全性定理（Gödel, 1931）"
    "和侯世达的自指怪圈（Hofstadter, 1979）放在同一张桌上。"
    "这三个框架各自深刻——但彼此隔绝。信息论不知道自指是免费的，"
    "不完全性定理不知道自指的经济成本，怪圈不知道它有一个运行着的实例。"
)
add_para(
    "本文提出GEME——一个同时实现了三条线索的运行系统。"
    "它不是一个神经网络，不是一个贝叶斯模型，不是一个符号系统。"
    "它是一个竞争性帧经济：帧在三条固定规则下竞争存活，无自由参数。"
)
add_para(
    "我们报告三个系统性发现：（1）自指的互信息趋近于零——香农-哥德尔桥；"
    "（2）第4层自指的行为相变——意识涌现的计算对应；"
    "（3）上层帧数收敛到6±2——与认知科学和社会网络的经验常数巧合。"
    "第2节描述模型，第3节报告实验，第4节讨论。"
)

# ==================== 2. 模型 ====================
doc.add_heading('2. GEME模型', level=1)

doc.add_heading('2.1 核心公理', level=2)
add_axiom(1, '竞争性合并',
    '输入向量到达时，与所有现有帧计算欧氏距离。若最近距离小于自适应阈值，'
    '输入合并入最接近帧，更新质心并增加权重；否则创建新帧。')
add_axiom(2, '自适应遗忘',
    '帧权重随时间衰减。低于存活阈值的帧被剪枝。衰减率依赖于帧合并历史。')
add_axiom(3, '自指观测',
    '系统定期从当前帧经济的聚合状态生成自观测向量，'
    '将其反馈到同一竞争合并过程，形成闭环。')

doc.add_heading('2.2 形式化定义', level=2)
add_para('帧 f 是一个五元组 (vec, w, age, m, sig)，其中：')
p = doc.add_paragraph('  vec ∈ ℝ²⁷（固定维度版本）；w ∈ ℝ⁺ 为权重；age ∈ ℕ 为步龄；m ∈ ℕ 为合并次数；sig 为可读标签。')
p.paragraph_format.first_line_indent = Pt(24)

add_formula('d(f₁, f₂) = ||vec₁ − vec₂||')
add_formula('w(f) = w₀ + Σ merge_event  − λ · age')
add_formula('阈值 θ = median{d(fᵢ, fⱼ)|∀i,j} · δ_adaptive')

doc.add_heading('2.3 分层动力学', level=2)
add_para(
    'L1（实体层）：δ-自适应合并将输入向量压缩为实体帧。'
    'L2（关联层）：滑动窗口检测帧间共现，生成──关联帧。'
    'L3（桥接层）：检测稳定关联模式，生成桥接帧——系统的第一个"自指"结构。'
    'L4（元认知层）：追踪L3桥接帧的权重变化率 d(w)/dt，'
    '生成关于"知识变化"的二阶自指帧。'
)

# ==================== 3. 实验发现 ====================
doc.add_heading('3. 实验发现', level=1)

doc.add_heading('3.1 香农-哥德尔桥：自指的资源中性', level=2)
add_para(
    '计算自指桥接帧 φ 与输入序列 X 之间的互信息 I(φ; X)。'
    '在3000+步运行中，I(φ; X) → 0，证明自指帧不携带输入的额外信息。'
)
add_formula('I(φ; X) = H(φ) − H(φ|X) → 0')
add_formula('C(X) = n · log₂|Σ|  （香农信道容量）')
add_formula('∴ I(φ; X) ≪ C(X)  ⇒  φ 不消耗信道容量')
add_para(
    '这等价于全同态加密（Gentry, 2009）在信息论层面的对偶：'
    'FHE证明计算可以在不访问内容的情况下进行；'
    '香农-哥德尔桥证明自指可以在不消耗信道容量的情况下发生。')
add_formula('差异 = 记忆容量为8时 ΔH = 0.162\n差异 = 记忆容量为32时 ΔH = 0.120\n差异 = 记忆容量为128时 ΔH = 0.041')
add_para('差异随1/N衰减，符合渐近自由性质。')

doc.add_heading('3.2 第4层相变', level=2)
add_para(
    'L4自指层存在一个临界复杂度阈值。低于阈值时，L4帧权重≈0；'
    '超过阈值时，L4桥接帧从0跳跃至356单位权重——一个清晰的相变。'
    '等价于人类元认知的现象：系统开始"知道自己在知道"。'
)
add_formula('L4桥接帧权重: 0  →  356  (在记忆容量=32时)')
add_formula('相变条件: 输入多样性 > 临界阈值')

doc.add_heading('3.3 魔数6：稳定吸引子', level=2)
add_para(
    '在记忆容量从8到52的扫描中，L4稳定帧数收敛到4-8，中心值约6。'
    '这个数字与三个独立领域的经验常数一致：'
)
items = [
    'Miller 7±2（工作记忆容量）：认知心理学70年经典结果',
    'Milgram 六度分割（社会网络直径）：任何两人之间平均6步',
    'Q + G ≈ PA（形式系统收敛）：哥德尔句子≈归纳法的经济等价物',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
add_formula('N_L4 ≈ log₂(N_total)  ≈ 6±2  （当 N_total = 32 时）')

# ==================== 4. 讨论 ====================
doc.add_heading('4. 讨论', level=1)

doc.add_heading('4.1 主要发现', level=2)
add_para(
    'GEME展示了：一个基于三条公理、无自由参数的极小帧经济，'
    '足以产生丰富的涌现结构。自指免费，第4层有相变，上层收敛到6——'
    '这些都不是预设的，是经济运行的结果。'
)

doc.add_heading('4.2 与认知架构的对话', level=2)
add_para(
    '自指的资源中性与Aaronson（2013, 2011）的计算复杂度进路一致。'
    'L4相变在结构上平行于Tononi的整合信息理论（Oizumi等，2014）中的意识阈值。'
    '竞争性帧经济与Friston（2010）的自由能原理共享"系统通过自适应更新最小化不确定性"的逻辑。'
)

doc.add_heading('4.3 统一性启示', level=2)
add_para(
    '数字6在心理学（Miller）、社会学（Milgram）和数学（Q+G≈PA）中的收敛，'
    '暗示了一个共同的信息经济原理在三个尺度的投影。'
    '这不是多学科巧合——是自指信道容量的自然上限。'
)

doc.add_heading('4.4 局限与未来方向', level=2)
add_para(
    'GEME目前十分抽象。输入空间是合成的27维向量。'
    '未来方向：（a）与LLM耦合实现PhasePrompt；'
    '（b）自然语言输入的多语言实验；'
    '（c）相变的数学分析；'
    '（d）多智能体社会模拟。'
)

doc.add_heading('4.5 结论', level=2)
add_para(
    'GEME是一个开源、可复现的计算棱镜。'
    '它不是一个认知理论——它是一个运行的系统。'
    '你可以打开终端，跑一次，看它自己涌现。'
)

# ==================== 致谢 ====================
doc.add_heading('致谢', level=1)
add_para(
    '感谢道格拉斯·R·侯世达教授的开创性工作——其关于自指与怪圈的思想是本研究的基石灵感，'
    '以及在最开始的鼓励。本研究由研究者独立完成，部分实验代码在AI辅助下完成，'
    '研究者对所有内容负责。代码与数据见 https://github.com/JackeyLGene/GEME。'
)

# ==================== 参考文献 ====================
doc.add_heading('参考文献', level=1)
refs_cn = [
    'Aaronson, S. (2011). Why philosophers should care about computational complexity.',
    'Aaronson, S. (2013). The ghost in the quantum Turing machine. arXiv:1306.0159.',
    'Friston, K. (2010). The free-energy principle. Nature Reviews Neuroscience.',
    'Gödel, K. (1931). Über formal unentscheidbare Sätze. Monatshefte für Mathematik.',
    'Hofstadter, D. R. (1979). Gödel, Escher, Bach. Basic Books.',
    'Miller, G. A. (1956). The magical number seven. Psychological Review.',
    'Milgram, S. (1967). The small world problem. Psychology Today.',
    'Oizumi, M. et al. (2014). From phenomenology to mechanisms of consciousness. PLoS Comp. Biol.',
    'Shannon, C. E. (1948). A mathematical theory of communication. Bell System Technical Journal.',
]
for ref in refs_cn:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)

# ==================== 补充材料 ====================
doc.add_page_break()
doc.add_heading('补充材料', level=1)

doc.add_heading('S1. 探索笔记：概念演化与科学范式对照', level=2)
add_para(
    'GEME开发过程观察到，帧经济的行为分化在结构上令人联想到不同的科学描述范式。'
    'L1帧的工作方式类似统计力学（离散粒子的聚合行为）；'
    'L2的共现关联类似热力学（时间方向性的涌现）；'
    'L3的桥接帧类似广义相对论（概念"质量"弯曲"认知时空"）；'
    'L4的自指类似量子测量（观测者与被观测系统的耦合）。'
    '这些对照是启发式的，标注于此以说明研究动机，不构成正式科学主张。')

doc.add_heading('S2. 泛化应用：汉谟拉比法典', level=2)
add_para(
    '将汉谟拉比法典的中文和英文译本分别输入两个GEME实例。'
    '尽管操作在完全不同的字符集上，L2关联拓扑结构收敛：'
    '中文产生27个关联帧，英文30个，帧权重的聚类分组在法律概念上一致'
    '（刑罚层次、损害赔偿、社会角色区分）。')

doc.add_heading('S3. 完整实验数据', level=2)
add_para('所有参数扫描、统计检验和控制实验的完整数据见 GitHub 仓库。')

doc.add_heading('S4. 复现指南', level=2)
add_para(
    '所有实验可通过 /standalone 目录下的脚本运行。依赖：Python 3.10+, torch, transformers, datasets。'
    '无需外部数据下载。见 README.md。')

# ==================== 保存 ====================
outpath = r'g:\GEME\preprint\GEME_中文提纲.docx'
doc.save(outpath)
copy_path = r'C:\Users\Administrator.DESKTOP-EM03IHL\Desktop\GEME_中文提纲.docx'
try:
    import shutil
    shutil.copy(outpath, copy_path)
    print(f'已保存到桌面')
except:
    pass
print(f'已保存: {outpath}')
print(f'大小: {os.path.getsize(outpath)/1024:.0f} KB')
