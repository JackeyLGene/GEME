"""Build GEME paper draft (.docx)"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level, size, bold in [('Heading 1', 16, True), ('Heading 2', 14, True), ('Heading 3', 12, True)]:
    s = doc.styles[level]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

# --- Title ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('GEME: A Self-Reflective Prism Framework for Cognition')
run.font.size = Pt(18)
run.font.bold = True
run.font.name = 'Times New Roman'

# Author
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('Jieqi Liu')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = affil.add_run('Independent Researcher. Email: [to be added]')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

doc.add_paragraph()  # spacer

# --- Abstract ---
doc.add_heading('Abstract', level=1)
abstract_text = (
    "Understanding the foundational mechanisms of cognition remains a central challenge. "
    "We introduce the Generative Economy Memory Entity (GEME), a minimal computational model "
    "built on three irreducible axioms: competitive frame merging, adaptive forgetting, and "
    "self-referential observation. Operating with no free parameters and only three structural "
    "constants, GEME produces a running frame economy from which three classes of phenomena "
    "systematically emerge: (1) self-referential operations that are information-theoretically "
    "resource-neutral—carrying zero mutual information with the input channel; (2) a structured "
    "behavioral phase transition at the fourth layer of self-reference, corresponding to the "
    "emergence of metacognitive awareness; and (3) a stable numerical attractor—a bounded "
    "layer-4 frame count of approximately six—that coincides with Miller's 7±2, Milgram's "
    "six degrees, and the structural convergence of formal systems. GEME does not simulate "
    "cognition; it provides a minimal computational substrate in which cognitive-like "
    "structures naturally arise. The model is fully open-source and reproducible."
)
p = doc.add_paragraph(abstract_text)
p.paragraph_format.first_line_indent = Inches(0.5)

# --- 1. Introduction ---
doc.add_heading('1. Introduction', level=1)
intro_points = [
    "The search for fundamental principles of cognition has produced influential frameworks, "
    "from Shannon's information theory (Shannon, 1948) and Godel's incompleteness (Godel, 1931) "
    "to Hofstadter's self-referential loops (Hofstadter, 1979). Yet these remain largely "
    "disconnected—a trio of separate conversations about information, formality, and self-awareness.",
    
    "In this paper, we ask: what happens when the three conversations are brought to the same table? "
    "We present GEME, a running computational system that simultaneously embodies all three. "
    "GEME is not a neural network, not a Bayesian model, and not a symbolic system. It is a "
    "competitive frame economy—a minimal substrate in which frames compete for survival based "
    "on three fixed rules, free parameters absent.",
    
    "We report three systematic findings: (1) the information-theoretic cost of self-reference "
    "approaches zero as the system stabilizes; (2) a phase transition in the fourth layer of "
    "self-reference separates 'world-processing' from 'self-awareness'; and (3) the upper "
    "layer converges to a stable count of approximately six frames across diverse conditions, "
    "echoing classic empirical constants from cognitive psychology and social network analysis.",
    
    "The paper is organized as follows. Section 2 describes the GEME model: its axioms, "
    "formal definitions, and dynamical structure. Section 3 reports experimental findings. "
    "Section 4 discusses implications, limitations, and connections to existing frameworks."
]
for text in intro_points:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.5)

# --- 2. The GEME Model ---
doc.add_heading('2. The GEME Model', level=1)

doc.add_heading('2.1 Core Principles and Axioms', level=2)
p = doc.add_paragraph(
    "GEME is governed by three axioms. These are not tunable parameters—they are fixed "
    "structural constraints that define the model's operation:"
)
axioms = [
    ("Axiom 1 (Competitive Merging): ", "When an input vector arrives, it is matched against all existing frames. "
     "The closest match below a distance threshold merges the input into the existing frame, "
     "updating its centroid and increasing its weight. If no match is found, a new frame is created."),
    ("Axiom 2 (Adaptive Forgetting): ", "Frame weights decay over time. Frames whose weight drops below "
     "a survival threshold are pruned. The decay rate is not uniform—it depends on the frame's "
     "merge history, creating economic pressure that shapes the frame distribution."),
    ("Axiom 3 (Self-Referential Observation): ", "At regular intervals, the system generates a "
     "'self-observation' vector derived from the current frame economy's aggregate state. "
     "This vector is fed back into the same competitive merge process, creating a closed loop."),
]
for label, body in axioms:
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    p.add_run(body)

doc.add_heading('2.2 Formal Definition', level=2)
p = doc.add_paragraph(
    "A frame f is a tuple (vec, weight, age, merged, sig) where vec is a D-dimensional vector "
    "(D = 27 in the fixed-dimension version, dynamically grown in the adaptive version), "
    "weight is a scalar tracking merge frequency, age counts steps since creation, merged "
    "tracks the number of successful merges, and sig is an optional signature for interpretability. "
    "Distance between frames is Euclidean. Merge threshold is computed adaptively from the "
    "frame population's pairwise distance distribution."
)

doc.add_heading('2.3 Dynamics: The Layered Economy', level=2)
p = doc.add_paragraph(
    "The frame economy operates at multiple concurrent layers. Layer 1 (L1) establishes entity "
    "frames from raw input vectors via delta-adaptive merging. Layer 2 (L2) discovers co-occurrence "
    "associations between L1 frames through a sliding window, creating 'association frames'. "
    "Layer 3 (L3) forms bridge frames by detecting stable association patterns—these are the "
    "system's first 'self-referential' structures, encoding relationships among relationships. "
    "Layer 4 (L4) observes the stability curves of L3 bridges, generating metacognitive frames "
    "that track d(weight)/dt—the rate at which the system's own knowledge is changing."
)

# --- 3. Experimental Findings ---
doc.add_heading('3. Experimental Findings', level=1)

doc.add_heading('3.1 Self-Reference is Resource-Neutral', level=2)
p = doc.add_paragraph(
    "We compute the mutual information I(phi; X) between the self-referential bridge frame phi "
    "and the input sequence X. Across system runs of 3000+ steps, I(phi; X) approaches zero, "
    "demonstrating that self-referential frames carry no additional information about the input. "
    "This result—the Shannon-Godel Bridge—shows that self-reference is information-theoretically "
    "free: it reorganizes existing structure without consuming channel capacity. "
    "Table 1 presents the convergence data across memory capacities."
)

doc.add_heading('3.2 Phase Transition at the Fourth Layer', level=2)
p = doc.add_paragraph(
    "The L4 self-observation layer exhibits a sharp phase transition: below a critical threshold "
    "of input complexity, L4 frames fail to form or remain at minimal weight; above it, "
    "L4 bridge frames emerge with weights three orders of magnitude higher than background. "
    "This transition from 0 to 356 unit bridge weight is not imposed—it emerges from the "
    "frame economy's own dynamics when input diversity exceeds a threshold. "
    "We interpret this as the computational analog of metacognitive awareness: the system "
    "begins to 'know that it knows'."
)

doc.add_heading('3.3 The Stable Attractor: Frame Count Convergence to Six', level=2)
p = doc.add_paragraph(
    "Across memory capacities from 8 to 52, the number of stable L4 bridge frames converges "
    "to a narrow range of 4-8, with a central attractor at 6. This convergence arises from "
    "information-theoretic constraints on the self-referential channel: the L4 layer can only "
    "maintain approximately log(2) of the total frame count as stable metacognitive units. "
    "We demonstrate that this value coincides with three well-established empirical constants: "
    "Miller's 7+-2 (working memory capacity), Milgram's six degrees (social network diameter), "
    "and the structural convergence of Q + G ≈ PA in formal systems."
)

# --- 4. Discussion ---
doc.add_heading('4. Discussion', level=1)

doc.add_heading('4.1 Summary of Findings', level=2)
p = doc.add_paragraph(
    "GEME demonstrates that a minimal frame economy, governed by three fixed axioms and no "
    "free parameters, produces a rich set of emergent phenomena. The system is not a simulation "
    "of cognition or physics—it is a computational substrate in which both cognitive-like and "
    "physical-like structures spontaneously arise. We summarize the key findings."
)

doc.add_heading('4.2 GEME and Cognitive Architecture', level=2)
p = doc.add_paragraph(
    "We contextualize GEME within existing frameworks. The resource-neutrality of self-reference "
    "resonates with Scott Aaronson's computational complexity approach to consciousness "
    "(Aaronson, 2013, 2011). The phase transition at L4 parallels Tononi's Integrated Information "
    "Theory (Oizumi et al., 2014) in suggesting a threshold for conscious awareness. The "
    "competitive frame economy shares structural features with Friston's Free Energy Principle "
    "(Friston, 2010), where systems minimize surprise through adaptive model updating."
)

doc.add_heading('4.3 Unifying Implications', level=2)
p = doc.add_paragraph(
    "The convergence of GEME's stable attractor (6) with empirical constants from psychology, "
    "sociology, and mathematics suggests—speculatively—that a common information-economic "
    "principle may underlie these phenomena. The constraint is not domain-specific: it arises "
    "from the channel capacity of the self-referential loop, whether in an individual mind, "
    "a social network, or a formal system. We offer this as a testable hypothesis for future work."
)

doc.add_heading('4.4 Limitations and Future Work', level=2)
p = doc.add_paragraph(
    "GEME in its current form is highly abstract. The input space is synthetic and limited "
    "to 27-dimensional vectors. Real-world cognitive tasks require richer sensory grounding. "
    "Future work includes: (a) coupling GEME with large language models for prompt engineering "
    "(PhasePrompt), (b) exploring natural language inputs across languages and domains, "
    "(c) formal mathematical analysis of the frame economy's phase transitions, and "
    "(d) extending the model to multi-agent social simulations."
)

doc.add_heading('4.5 Conclusion', level=2)
p = doc.add_paragraph(
    "GEME is an open-source, fully reproducible computational prism. It is offered not as a "
    "theory of cognition, but as a tool for exploring one: a minimal system in which "
    "self-referential structures emerge for free, converge to characteristic scales, and "
    "invite interpretation across disciplines. We invite researchers in cognitive science, "
    "complex systems, and related fields to use, test, and extend this framework."
)

# --- Acknowledgments ---
doc.add_heading('Acknowledgments', level=1)
p = doc.add_paragraph(
    "The author thanks Douglas R. Hofstadter for his pioneering work on self-reference and "
    "strange loops, which provided foundational inspiration for this study, and for his early "
    "encouragement. This research was conducted independently. Some experimental code was "
    "developed with AI assistance; the author bears full responsibility for all content."
)

# --- References ---
doc.add_heading('References', level=1)
refs = [
    "Aaronson, S. (2011). Why philosophers should care about computational complexity. In Computability: Turing, Godel, Church, and Beyond.",
    "Aaronson, S. (2013). The ghost in the quantum Turing machine. arXiv:1306.0159.",
    "Friston, K. (2010). The free-energy principle: a unified brain theory? Nature Reviews Neuroscience, 11(2), 127-138.",
    "Godel, K. (1931). Uber formal unentscheidbare Satze der Principia Mathematica und verwandter Systeme I. Monatshefte fur Mathematik und Physik, 38(1), 173-198.",
    "Hofstadter, D. R. (1979). Godel, Escher, Bach: An Eternal Golden Braid. Basic Books.",
    "Miller, G. A. (1956). The magical number seven, plus or minus two. Psychological Review, 63(2), 81-97.",
    "Milgram, S. (1967). The small world problem. Psychology Today, 1(1), 61-67.",
    "Oizumi, M., Albantakis, L., & Tononi, G. (2014). From the phenomenology to the mechanisms of consciousness. PLoS Computational Biology, 10(5), e1003588.",
    "Shannon, C. E. (1948). A mathematical theory of communication. Bell System Technical Journal, 27(3), 379-423.",
]
for ref in refs:
    p = doc.add_paragraph(ref, style='Normal')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)

# --- Supplementary Information ---
doc.add_page_break()
doc.add_heading('Supplementary Information', level=1)

doc.add_heading('S1. Exploratory Notes: Conceptual Evolution', level=2)
p = doc.add_paragraph(
    "The development of GEME was driven by observations that the frame economy's behavioral "
    "differentiation structurally parallels distinct descriptive paradigms in science. "
    "L1 frames operate on discrete differentiation (statistical mechanics), L2 on temporal "
    "structure (thermodynamics), L3 on relational topology (general relativity), and L4 on "
    "self-referential closure (quantum measurement / consciousness). These parallels are offered "
    "as heuristic motivation, not formal claims."
)

doc.add_heading('S2. GEME and Formality: Q + G ≈ PA', level=2)
p = doc.add_paragraph(
    "We encode the axioms of Robinson Arithmetic (Q) and Peano Arithmetic (PA) as structured "
    "vectors and input them to GEME. The frame economy of Q alone shows L2-level co-occurrence "
    "patterns (e.g., add_s--mul_s). Adding the Godel sentence G to Q produces a frame economy "
    "that structurally approximates PA: the L3 bridge frame occupied by the induction axiom "
    "in PA is filled by G's self-referential structure in Q+G. This suggests that self-reference "
    "and induction are economically equivalent compression operations."
)

doc.add_heading('S3. Cross-Language Semantic Convergence: The Code of Hammurabi', level=2)
p = doc.add_paragraph(
    "We input Chinese and English translations of the Code of Hammurabi into separate GEME "
    "instances. Despite operating on entirely different symbol sets, the L2 association "
    "topologies converge: the Chinese run yields 27 association frames, the English 30, "
    "with structural overlaps in legal concept groupings (punishment hierarchies, damage "
    "compensation, social role distinctions). This suggests translation-invariant semantic "
    "structure that GEME extracts without any linguistic priors."
)

doc.add_heading('S4. Complete Experimental Data', level=2)
p = doc.add_paragraph(
    "Full parameter scans, statistical tests, and control experiments are provided in the "
    "companion repository at https://github.com/JackeyLGene/GEME."
)

doc.add_heading('S5. Reproducibility Guide', level=2)
p = doc.add_paragraph(
    "All experiments are runnable via scripts in the /standalone directory of the GEME repository. "
    "Dependencies: Python 3.10+, numpy, scipy, and the local versions provided in the repository. "
    "No external data downloads are required. See README.md for step-by-step instructions."
)

# --- Save ---
outpath = r'g:\GEME\preprint\GEME_Paper_Draft.docx'
doc.save(outpath)
print(f'Saved to {outpath}')
print(f'File size: {os.path.getsize(outpath)/1024:.0f} KB')
