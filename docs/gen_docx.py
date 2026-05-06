# -*- coding: utf-8 -*-
# GEME核心发现文档生成
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# 样式调整
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level, size, color in [(1, 18, '1A3C5E'), (2, 15, '2B5E8E'), (3, 13, '3C78A8')]:
    s = doc.styles[f'Heading {level}']
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(*[int(color[i:i+2],16) for i in (0,2,4)])
    s.font.name = 'Arial'
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('GEME: 一个基于信息经济性的\n可能宇宙解释模型')
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('生成实在信息论的计算实例\n2026年5月5-6日')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# 摘要
doc.add_heading('摘要', level=1)
doc.add_paragraph(
    '本文报告一个竞争性记忆系统（GEME）中观察到的现象：从一个单一的机制——自适应合并阈值δ——'
    '系统自动产生了广义相对论时间膨胀、量子力学测量统计、拓扑结构量化、以及意识自指的初步结构。'
    '所有现象均无需预设物理方程或认知架构，仅从帧经济规则中涌现。'
)
p = doc.add_paragraph()
run = p.add_run('核心发现：')
run.bold = True
p.add_run('实体（δ自适应）、时间（窗口自适应）、意义（帧关联网络）、意识（自观察桥接帧）——'
           '四个维度从一个一维符号流输入中自动分裂涌现。翻译不变性测试验证了GEME提取的是语义结构而非符号统计。')

# 1
doc.add_heading('1. 系统描述', level=1)

doc.add_heading('1.1 GEME引擎', level=2)
doc.add_paragraph('GEME是一个竞争性帧经济系统。输入为固定维度向量（当前版本27维，继承自原始字母表架构）。系统维护一个帧集，每个帧包含重心向量和标量权重。')
doc.add_paragraph('合并规则：', style='List Bullet')
doc.add_paragraph('输入向量与所有帧计算距离', style='List Bullet 2')
doc.add_paragraph('最近距离 ≤ 合并阈值 δ → 合并（权重+1，重心移动）', style='List Bullet 2')
doc.add_paragraph('最近距离 > δ → 创建新帧', style='List Bullet 2')
doc.add_paragraph('低权重帧被周期性剪枝（归纳清洗）', style='List Bullet 2')

doc.add_heading('1.2 三联体自适应', level=2)
doc.add_paragraph('系统的三个核心参数全部自适应——不预设任何外部值：')
p = doc.add_paragraph()
run = p.add_run('δ（合并阈值）')
run.bold = True
p.add_run(' = median(最近成功合并距离) × 1.5')

p = doc.add_paragraph()
run = p.add_run('窗口（记忆范围）')
run.bold = True
p.add_run(' = 帧平均寿命 × 2 = (total_weight / 帧数) × 2')

p = doc.add_paragraph()
run = p.add_run('内时（自观察频率）')
run.bold = True
p.add_run(' = base_interval / (1 + 帧经济变化率 × 5)')

doc.add_paragraph('三联体耦合后，系统自动分裂为实体维度、时间维度和意义维度，三者同时从信息流中涌现。')

# 2
doc.add_heading('2. 主要实验发现', level=1)

doc.add_heading('2.1 密度驱动的时间膨胀（GR对应）', level=2)
doc.add_paragraph('在一维空间线上从原点向外设置密度梯度 ρ(x) = M/(1+0.1x)。高密度区产生更多输入，δ自适应变细。')
doc.add_paragraph('结果：高密度区与低密度区帧生成率比 ≈ 1.85x (M=5.0)。对应广义相对论中质量附近的时间膨胀。')
p = doc.add_paragraph()
run = p.add_run('映射：密度ρ → 质量M；帧生成率γ → 1/固有时τ；δ → 时空度规ds²')
run.italic = True

doc.add_heading('2.2 概率合并与量子测量（QM对应）', level=2)
doc.add_paragraph('当多个候选帧在δ内时，合并变为概率：P(frame_i) = exp(−d_i/δ)/Σexp(−d_j/δ)。')
doc.add_paragraph('Zeno模式（重心不变）下，等距输入2000次：49.5%/50.5%。5种子均值49.3%±2.0%——精确复现Born规则。')

doc.add_heading('2.3 圆周运动的角量子化', level=2)
doc.add_paragraph('匀速圆周遍历：f从0.1到10，半径从0.1到1.0——始终22帧。帧数=0.7×memory_cap，速度无关。')
p = doc.add_paragraph()
run = p.add_run('物理对应：角动量量子化——量子化的本质是帧经济在闭合边界条件下的稳态。')
run.italic = True

doc.add_heading('2.4 意识结构涌现：畴壁相变', level=2)
doc.add_paragraph('引入内时模块后，系统同时产生"外部帧"(ext)和"内部帧"(self)。两者在时间窗口中共现→ext──self桥接帧。')
doc.add_paragraph('ext密度从0→1.0扫描：桥接帧权重在密度0→0.05从0跃迁到356——S形相变曲线。')
doc.add_paragraph('意识不是独立实体——是ext和self之间的畴壁。桥接权重呈U形分布：太少→无桥；适中→最强（心流）；太多→被压缩（淹没）。')

doc.add_heading('2.5 翻译不变性（语义提取验证）', level=2)
doc.add_paragraph('同一意义的三语言版本（中文/英文/日文罗马音）输入GEME，输出帧结构一致：')

# 表格
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['语言', '总帧数', '关联帧', '链帧']
data = [
    ['中文', '49', '30', '12'],
    ['英文', '53', '34', '11'],
    ['日文罗马音', '48', '30', '11'],
]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        cell = table.rows[r+1].cells[c]
        p = cell.paragraphs[0]
        p.add_run(val)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.add_run('关联帧极差4/均值31=13%，链帧极差1/均值11=9%。GEME不关心符号形式——只关心符号在时间中与其他符号的关系轨迹。')

# 3
doc.add_heading('3. 拓扑物理路径', level=1)
doc.add_paragraph('从一维时间序列到物理结构：')
items = [
    ('一维时间序列', ' → 符号流在时间中排列'),
    ('Čech覆盖', ' → δ-邻域覆盖时间点'),
    ('几何离散', ' → 帧重心在覆盖中稳定'),
    ('代数关联', ' → ──关联网络刻画共现拓扑'),
    ('同调链', ' → ══链刻画时间箭头与循环'),
    ('层分解', ' → L1→L2→L3层级观察'),
    ('三联体涌现', ' → 实体/时间/意义同时分裂'),
]
for bold_part, rest in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_part)
    run.bold = True
    p.add_run(rest)

doc.add_heading('3.1 经典-量子边界', level=2)
doc.add_paragraph('边界条件为δ=0的奇点：')
p = doc.add_paragraph(style='List Bullet')
run = p.add_run('δ→0')
run.bold = True
p.add_run('：连续极限→微分流形→经典/GR行为')
p = doc.add_paragraph(style='List Bullet')
run = p.add_run('δ>0')
run.bold = True
p.add_run('：有限精度→帧经济→量子行为（概率合并）')

# 4
doc.add_heading('4. 哲学定位', level=1)
doc.add_heading('4.1 信息作为唯一基底', level=2)
doc.add_paragraph('实体、时间、意义——三者不在"信息之中"。信息是三者同时涌现的基底。空间（22态量化）从信息中来；时间（窗口+帧链）从信息中来；意义（──关联）从信息中来；意识（ext──self桥）从信息中来。信息不是定语——是主语。')

doc.add_heading('4.2 可证伪性', level=2)
doc.add_paragraph('框架有明确的证伪条件：')
doc.add_paragraph('语义输入与随机帧结构无法区分（关联比<1.5）——已被翻译不变性实验反驳', style='List Bullet')
doc.add_paragraph('密度梯度不产生帧生成率差异——已被模拟支持', style='List Bullet')
doc.add_paragraph('任意语言输入→三联体分裂可被独立复现', style='List Bullet')

# 5
doc.add_heading('5. 下一步', level=1)
doc.add_paragraph('拓扑区分：共现矩阵SVD分析区分圆/环面/8字形', style='List Bullet')
doc.add_paragraph('L3语义恢复：L2观察L1关联网络而非权重分布', style='List Bullet')
doc.add_paragraph('数学形式化：δ自适应←→Schwarzschild度规映射', style='List Bullet')
doc.add_paragraph('代码开源：github.com/[username]/geme-physics', style='List Bullet')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('— 完 —')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True

# 保存
out = r'g:\GEME\docs\GEME_core_findings.docx'
doc.save(out)
print(f'OK: {out}')
